
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import base64
from docx import Document
from io import BytesIO
import os
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
from django.conf import settings
from ..utils import parse_datetime
from docx.shared import Inches

@csrf_exempt
def doc_competitive(request):
    channel_icons = request.session.get('channel_icons', [])
    channel_names = request.session.get('channel_names', [])
    top_videos = request.session.get('top_videos', [])
    durations = request.session.get('durations', [])
    output_data = {
        'average_likes': request.session['average_likes'],
        'average_views': request.session['average_views'],
        'subs': request.session['subs'],
        'channel_names': channel_names,
        'durations': durations,
        'mostUsedCategories': request.session.get('mostUsedCategories', []),
        'topTags': request.session.get('topTags', []),
    }
    json_data = json.dumps(output_data)
    

    # Zip the lists together in the view
    channels = zip(channel_icons, channel_names)
    channels_tags = zip(request.session.get('topTags', []), channel_names)
    context = {
        'channels': channels,
        'json_data': json_data,
        'top_likes_channel': request.session['top_likes_channel'],
        'top_views_channel': request.session['top_views_channel'],
        'top_subs_channel': request.session['top_subs_channel'],
        'type': request.session['type'],
        'top_videos': top_videos,
        'output_data': output_data,
        'channels_tags': channels_tags
        
    }
    if request.method == 'POST':
        doc = Document()
        doc.add_heading('Compeitive Analysis', 0)
        doc.add_paragraph('Our youtube competitive analysis will provide you with customizable dataset, statistics, graphs and interpretaions to make your work with data easier.')
        
        doc.add_heading('Analysed playlist:', level=1)
        for name in channel_names:
            doc.add_paragraph(name, style='ListBullet')

        doc.add_heading('Top videos from each channel:', level=1)
        for video in top_videos:
            # Heading for each video title
            doc.add_heading(video['title'], level=2)

            # Bold description label
            desc_para = doc.add_paragraph()
            desc_para.add_run('Description: ').bold = True
            desc_para.add_run(video['description'])
            
            # Bold statistics label
            stats_para = doc.add_paragraph()
            stats_para.add_run('Statistics: ').bold = True
            stats_para.add_run(f"{video['viewsCount']} views, {video['likesCount']} likes, {video['duration']} minutes")

            # Comments as bullet points
            if video['topComments']:  # Check if there are comments
                comments_para = doc.add_paragraph()
                comments_para.add_run('Top Comments:').bold = True
                for comment in video['topComments']:
                    doc.add_paragraph(comment, style='ListBullet')
                    
        doc.add_heading('Get more insights with graphs:', level=1)
        
        data = json.loads(request.body)
        # The variable here should match what's sent from the frontend
        imgs_data = data['imgData']  # This should match the key in the JSON sent from the frontend
        for img_data in imgs_data:
            # Decode the base64 image
            img_data = base64.b64decode(img_data.split(',')[1])
            image_stream = BytesIO(img_data)
            # Add the image to the Word document
            doc.add_picture(image_stream)

        # Save the document
        output_dir = os.path.join(settings.MEDIA_ROOT, 'documents')
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        file_path = os.path.join(output_dir, "competitive_analysis.docx")
        doc.save(file_path)

        return JsonResponse({'message': 'Document created successfully'})

    return JsonResponse({'error': 'Invalid request'}, status=400)


@csrf_exempt
def doc_channel(request):
    
    channel_data= {
        
    
    'top_5_videos': request.session['top_5_videos'],
    'worst_5_videos': request.session['worst_5_videos'],
    'uniqueTags': request.session['uniqueTags'],
    'all_playlists_dict': request.session['all_playlists_dict'],         
    'title': request.session['title'],
    'description': request.session['description'],
    'thumbnail': request.session['thumbnail'],
    'videoCount': request.session['videoCount'],
    'totalViews': request.session['totalViews'],
    'totalLikes': request.session['totalLikes'],
    'totalComments': request.session['totalComments'],
    'average_duration': request.session['average_duration'],
    }
    
    
              
    
    if request.method == 'POST':
        doc = Document()
        doc.add_heading('Channel Analysis', 0)
        doc.add_paragraph('Our channel analysis will give you an overview over the channel, what does influence its performance, and a closer look on its top and worst performing videos...')
        
        doc.add_heading('Get An Overview', level=1)
        doc.add_heading(channel_data['title'], level=2)
        
        doc.add_heading('Discreption', level=3)
        doc.add_paragraph(channel_data['description'])
        doc.add_heading('Word tags', level=3)
        tags_str = ', '.join(channel_data['uniqueTags'])
        doc.add_paragraph(tags_str)
        doc.add_heading('Statistics', level=3)
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Video Count: ').bold = True
        stats_para.add_run(str(channel_data['videoCount']))  # Convert integer to string
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Total Views: ').bold = True
        stats_para.add_run(str(channel_data['totalViews']))  # Convert integer to string
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Total likes: ').bold = True
        stats_para.add_run(str(channel_data['totalLikes']))  # Convert integer to string
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Comments Count: ').bold = True
        stats_para.add_run(str(channel_data['totalComments']))  # Convert integer to string
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Videos Average Duration: ').bold = True
        stats_para.add_run(str(channel_data['average_duration']))  # Convert integer to string

        doc.add_heading('Top and worst videos', level=1)
        
        doc.add_heading('Top videos', level=2)
        doc.add_heading('Top videos info', level=3)
        for video in channel_data['top_5_videos']:
            # Heading for each video title
            doc.add_heading(video['title'], level=4)

            # Bold description label
            desc_para = doc.add_paragraph()
            desc_para.add_run('Description: ').bold = True
            desc_para.add_run(video['description'])
            
            # Bold statistics label
            stats_para = doc.add_paragraph()
            stats_para.add_run('Statistics: ').bold = True
            stats_para.add_run(f"{video['viewsCount']} views, {video['likesCount']} likes, {video['duration']} minutes")
        
        
        if('top_5_comments_analysis_dist' in request.session): 
            top_5_comments_analysis_dist = request.session['top_5_comments_analysis_dist']
            top_5_comments = request.session['top_5_comments']

            doc.add_heading('Top videos Comments and Sentiment', level=3)
            for emotion, comment in top_5_comments.items():
                stats_para = doc.add_paragraph(style='ListBullet')
                stats_para.add_run(f"{emotion} :").bold = True
                stats_para.add_run(comment)
            
            
            
        doc.add_heading('Worst videos', level=2)
        doc.add_heading('Worst videos info', level=3)
        for video in channel_data['worst_5_videos']:
            # Heading for each video title
            doc.add_heading(video['title'], level=4)

            # Bold description label
            desc_para = doc.add_paragraph()
            desc_para.add_run('Description: ').bold = True
            desc_para.add_run(video['description'])
            
            # Bold statistics label
            stats_para = doc.add_paragraph()
            stats_para.add_run('Statistics: ').bold = True
            stats_para.add_run(f"{video['viewsCount']} views, {video['likesCount']} likes, {video['duration']} minutes")

        if('worst_5_comments_analysis_dist' in request.session): 
            worst_5_comments_analysis_dist = request.session['worst_5_comments_analysis_dist']
            worst_5_comments = request.session['worst_5_comments']
        
            doc.add_heading('Worst videos Comments and Sentiment', level=3)
            for emotion, comment in worst_5_comments.items():
                stats_para = doc.add_paragraph(style='ListBullet')
                stats_para.add_run(f"{emotion} :").bold = True
                stats_para.add_run(comment)

        
        doc.add_heading('Get more insights with graphs:', level=1)
        
        data = json.loads(request.body)
        # The variable here should match what's sent from the frontend
        imgs_data = data['imgData']  # This should match the key in the JSON sent from the frontend
        for img_data in imgs_data:
            # Decode the base64 image
            img_data = base64.b64decode(img_data.split(',')[1])
            image_stream = BytesIO(img_data)
            # Add the image to the Word document
            doc.add_picture(image_stream)

        # Save the document
        output_dir = os.path.join(settings.MEDIA_ROOT, 'documents')
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        file_path = os.path.join(output_dir, "channel_analysis.docx")
        doc.save(file_path)

        return JsonResponse({'message': 'Document created successfully'})

    return JsonResponse({'error': 'Invalid request'}, status=400)



@csrf_exempt
def doc_playlist(request):
    
    channel_data= {
        
    
    'top_5_videos': request.session['top_5_videos'],
    'worst_5_videos': request.session['worst_5_videos'],
    'uniqueTags': request.session['uniqueTags'],
    'all_playlists_dict': request.session['all_playlists_dict'],         
    'title': request.session['title'],
    'description': request.session['description'],
    'thumbnail': request.session['thumbnail'],
    'videoCount': request.session['videoCount'],
    'totalViews': request.session['totalViews'],
    'totalLikes': request.session['totalLikes'],
    'totalComments': request.session['totalComments'],
    'average_duration': request.session['average_duration'],
    }
              
    
    if request.method == 'POST':
        doc = Document()
        doc.add_heading('Playlist Analysis', 0)
        doc.add_paragraph('Our playlist analysis will give you an overview over the playlist, what does influence its performance, and a closer look on its top and worst performing videos...')
        
        doc.add_heading('Get An Overview', level=1)
        doc.add_heading(channel_data['title'], level=2)
        
        doc.add_heading('Discreption', level=3)
        doc.add_paragraph(channel_data['description'])
        doc.add_heading('Word tags', level=3)
        tags_str = ', '.join(channel_data['uniqueTags'])
        doc.add_paragraph(tags_str)
        doc.add_heading('Statistics', level=3)
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Video Count: ').bold = True
        stats_para.add_run(str(channel_data['videoCount']))  # Convert integer to string
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Total Views: ').bold = True
        stats_para.add_run(str(channel_data['totalViews']))  # Convert integer to string
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Total likes: ').bold = True
        stats_para.add_run(str(channel_data['totalLikes']))  # Convert integer to string
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Comments Count: ').bold = True
        stats_para.add_run(str(channel_data['totalComments']))  # Convert integer to string
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Videos Average Duration: ').bold = True
        stats_para.add_run(str(channel_data['average_duration']))  # Convert integer to string

        doc.add_heading('Top and worst videos', level=1)
        
        doc.add_heading('Top videos', level=2)
        doc.add_heading('Top videos info', level=3)
        for video in channel_data['top_5_videos']:
            # Heading for each video title
            doc.add_heading(video['title'], level=4)

            # Bold description label
            desc_para = doc.add_paragraph()
            desc_para.add_run('Description: ').bold = True
            desc_para.add_run(video['description'])
            
            # Bold statistics label
            stats_para = doc.add_paragraph()
            stats_para.add_run('Statistics: ').bold = True
            stats_para.add_run(f"{video['viewsCount']} views, {video['likesCount']} likes, {video['duration']} minutes")

        if('top_5_comments_analysis_dist' in request.session): 
            top_5_comments_analysis_dist = request.session['top_5_comments_analysis_dist']
            top_5_comments = request.session['top_5_comments']

            doc.add_heading('Top videos Comments and Sentiment', level=3)
            for emotion, comment in top_5_comments.items():
                stats_para = doc.add_paragraph(style='ListBullet')
                stats_para.add_run(f"{emotion} :").bold = True
                stats_para.add_run(comment)

            
            
            
        doc.add_heading('Worst videos', level=2)
        doc.add_heading('Worst videos info', level=3)
        for video in channel_data['worst_5_videos']:
            # Heading for each video title
            doc.add_heading(video['title'], level=4)

            # Bold description label
            desc_para = doc.add_paragraph()
            desc_para.add_run('Description: ').bold = True
            desc_para.add_run(video['description'])
            
            # Bold statistics label
            stats_para = doc.add_paragraph()
            stats_para.add_run('Statistics: ').bold = True
            stats_para.add_run(f"{video['viewsCount']} views, {video['likesCount']} likes, {video['duration']} minutes")

        if('worst_5_comments_analysis_dist' in request.session): 
            worst_5_comments_analysis_dist = request.session['worst_5_comments_analysis_dist']
            worst_5_comments = request.session['worst_5_comments']
        
            doc.add_heading('Worst videos Comments and Sentiment', level=3)
            for emotion, comment in worst_5_comments.items():
                stats_para = doc.add_paragraph(style='ListBullet')
                stats_para.add_run(f"{emotion} :").bold = True
                stats_para.add_run(comment)
        
        doc.add_heading('Get more insights with graphs:', level=1)
        
        data = json.loads(request.body)
        # The variable here should match what's sent from the frontend
        imgs_data = data['imgData']  # This should match the key in the JSON sent from the frontend
        for img_data in imgs_data:
            # Decode the base64 image
            img_data = base64.b64decode(img_data.split(',')[1])
            image_stream = BytesIO(img_data)
            # Add the image to the Word document
            doc.add_picture(image_stream)

        # Save the document
        output_dir = os.path.join(settings.MEDIA_ROOT, 'documents')
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        file_path = os.path.join(output_dir, "playlist_analysis.docx")
        doc.save(file_path)

        return JsonResponse({'message': 'Document created successfully'})

    return JsonResponse({'error': 'Invalid request'}, status=400)



@csrf_exempt
def doc_topic(request):
    
    top_5_videos = request.session['top_5_videos']
    # Assuming 'top_5_videos' is a list of dictionaries containing video information
    for video in top_5_videos:
        # Parse the 'publishedAt' date string to a datetime object
        date_obj = parse_datetime(video['publishedAt'])

        # Reformat the date to 'YYYY MMM DD' and update the video dictionary
        video['publishedAt'] = date_obj.strftime('%Y %b %d')

            
    context= {
        'top_5_videos': top_5_videos,
        'channels_dict': request.session['channels_dict'],
        'videos_dict': request.session['videos_dict'],
        'top_5_comments': request.session['top_5_comments'],
              }
              
    
    if request.method == 'POST':
        doc = Document()
        doc.add_heading('Topic Analysis', 0)
        doc.add_paragraph('Check out the metrics and insights to dive into the trends and discussions around your chosen subject.')
                
        
        doc.add_heading('Top channels', level=1)
        for channel in context['channels_dict']:
            # Heading for each video title
            doc.add_heading(channel['Name'], level=2)

            

            doc.add_heading('Statistics', level=3)
            stats_para = doc.add_paragraph(style='ListBullet')
            stats_para.add_run('Video Count: ').bold = True
            stats_para.add_run(str(channel['Video count']))  # Convert integer to string
            stats_para = doc.add_paragraph(style='ListBullet')
            stats_para.add_run('Views average: ').bold = True
            stats_para.add_run(str(channel['TotalViews']))  # Convert integer to string
            stats_para = doc.add_paragraph(style='ListBullet')
            stats_para.add_run('likes average: ').bold = True
            stats_para.add_run(str(channel['TotalLikes']))  # Convert integer to string
            stats_para = doc.add_paragraph(style='ListBullet')
            stats_para.add_run('Subscriber count: ').bold = True
            stats_para.add_run(str(channel['Subscriber count']))  # Convert integer to string
            stats_para = doc.add_paragraph(style='ListBullet')
            stats_para.add_run('Playlist count: ').bold = True
            stats_para.add_run(str(channel['Playlist count']))  # Convert integer to string
            
            # Bold description label
            desc_para = doc.add_paragraph()
            
            
        
        doc.add_heading('Top videos', level=1)
        for video in context['top_5_videos']:
            # Heading for each video title
            doc.add_heading(video['title'], level=2)

            
            # Bold statistics label
            stats_para = doc.add_paragraph()
            stats_para.add_run('Statistics: ').bold = True
            stats_para.add_run(f"{video['viewsCount']} views, {video['likesCount']} likes, {video['duration']} minutes")
            # Bold description label
            desc_para = doc.add_paragraph()
            
            
            
        doc.add_heading('Top videos Comments and Sentiment', level=1)
        for emotion, comment in context['top_5_comments'].items():
            stats_para = doc.add_paragraph(style='ListBullet')
            stats_para.add_run(f"{emotion} :").bold = True
            stats_para.add_run(comment)
            
            
    
        

        
        doc.add_heading('Get more insights with graphs:', level=1)
        
        data = json.loads(request.body)
        # The variable here should match what's sent from the frontend
        imgs_data = data['imgData']  # This should match the key in the JSON sent from the frontend
        for img_data in imgs_data:
            # Decode the base64 image
            img_data = base64.b64decode(img_data.split(',')[1])
            image_stream = BytesIO(img_data)
            # Add the image to the Word document
            doc.add_picture(image_stream)

        # Save the document
        output_dir = os.path.join(settings.MEDIA_ROOT, 'documents')
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        file_path = os.path.join(output_dir, "topic_analysis.docx")
        doc.save(file_path)

        return JsonResponse({'message': 'Document created successfully'})

    return JsonResponse({'error': 'Invalid request'}, status=400)
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Pt
from docx.text.run import Run
import docx
def add_hyperlink(paragraph, url, text, color, underline, heading=False):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    if not underline:
        u = OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    if heading:
        run = Run(new_run, paragraph)
        run.font.size = Pt(16)  # Adjust the size as per your heading style
        run.font.bold = True

    return hyperlink


def def_retrive(request):
              
        doc = Document()
        doc.add_heading('Videos Retriving', 0)
        doc.add_paragraph('Explore the collection below to discover content related to your search. Use these insights to enhance your understanding, create content, or simply enjoy the diversity of videos available on your topic of interest.')
                
        related_videos_full_dict = request.session['related_videos_full_dict']
        
        doc.add_heading('List of Videos', level=1)
        for video in related_videos_full_dict:
            p = doc.add_paragraph()
            add_hyperlink(p, video['URL'], video['Title'], '0000FF', False, heading=True)
            stats_para = doc.add_paragraph()
            stats_para.add_run('Channel: ').bold = True
            stats_para.add_run(f"{video['Channel']}")
            stats_para = doc.add_paragraph()
            stats_para.add_run('Statistics: ').bold = True
            stats_para.add_run(f"{video['Views']} views, {video['Likes']} likes, {video['Comments']} comments, {video['Duration']} minutes")
            stats_para = doc.add_paragraph()
            stats_para.add_run('Category: ').bold = True
            stats_para.add_run(f"{video['Category']}")
            stats_para = doc.add_paragraph()
            stats_para.add_run('Top Tags: ').bold = True
            stats_para.add_run(f"{video['Tags'][:5]}")
    
        # Save the document
        output_dir = os.path.join(settings.MEDIA_ROOT, 'documents')
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        file_path = os.path.join(output_dir, "video_retriving.docx")
        doc.save(file_path)



@csrf_exempt
def doc_post(request):
    
    post_data= {
        
    
    'publishedAt': request.session['publishedAt'],
    'owner': request.session['owner'],
    'top_keywords': request.session['top_keywords'],
    'MediaCount': request.session['MediaCount'],         
    'LikeCount': request.session['LikeCount'],
    'CommentCount': request.session['CommentCount'],
    'comment_sentiments': request.session['comment_sentiments'],
    'caption': request.session['caption'],

}

    
    
              
    
    if request.method == 'POST':
        doc = Document()
        doc.add_heading('Post Analysis', 0)
        doc.add_paragraph('Our post analysis will give you an overview over the post, what does influence its performance, and a closer look on its top comments...')
        
        doc.add_heading('Get An Overview to the post information', level=1)

        data = json.loads(request.body)
        
        imgs_data = data['chartData'] #key of the dictonary of the post pic 
        thumbnialData = data['thumbnailData']
        
        for thumbnial in thumbnialData:
            # Decode the base64 image
            thumbnial = base64.b64decode(thumbnial.split(',')[1])
            image_stream = BytesIO(thumbnial)
            # Add the image to the Word document
            doc.add_picture(image_stream, width=Inches(4), height=Inches(4))
            
        
        
        doc.add_heading('post caption', level=3)
        doc.add_paragraph(post_data['caption'])


        doc.add_heading('owner', level=3)
        doc.add_paragraph(post_data['owner'])
        
        doc.add_heading('published At', level=3)
        doc.add_paragraph(post_data['publishedAt'])

        
        doc.add_heading('top keywords', level=3)
        keyWords_str = ', '.join(post_data['top_keywords'])
        doc.add_paragraph(keyWords_str)




        doc.add_heading('Top post Comments for each Sentiment', level=3)
        for emotion, comment in post_data['comment_sentiments'].items():
                stats_para = doc.add_paragraph(style='ListBullet')
                stats_para.add_run(f"{emotion} :").bold = True
                stats_para.add_run(comment)


        doc.add_heading('Statistics', level=3)
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Like Count: ').bold = True
        stats_para.add_run(str(post_data['LikeCount']))  

        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Comment Count: ').bold = True
        stats_para.add_run(str(post_data['CommentCount']))  

        doc.add_heading('Get more insights with graphs:', level=1)

        for img_data in imgs_data:
            # Decode the base64 image
            img_data = base64.b64decode(img_data.split(',')[1])
            image_stream = BytesIO(img_data)
            # Add the image to the Word document
            doc.add_picture(image_stream)

# Save the document
        output_dir = os.path.join(settings.MEDIA_ROOT, 'documents')
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        file_path = os.path.join(output_dir, "post_analysis.docx")
        doc.save(file_path)

        return JsonResponse({'message': 'Document created successfully'})

    return JsonResponse({'error': 'Invalid request'}, status=400)