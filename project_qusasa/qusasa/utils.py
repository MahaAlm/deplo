import pandas as pd
from datetime import timedelta
import re
from collections import Counter
from transformers import pipeline
import torch
from transformers import pipeline, BertTokenizer, AutoModelForSequenceClassification
import torch
from .youtube_api import get_youtube_client
from googleapiclient.errors import HttpError

#competitive analysis utils

def searchByQuery(youtube, keyword, Type, orderBy='relevance', regionCode='', language=''):
    # Prepare the request with the parameters
    Max = 4
    
    request_parameters = {
        "part": "snippet",
        "order": orderBy,
        "q": keyword,
        "type": Type,
        "maxResults": Max
    }
    
    if(regionCode != ''):
        request_parameters["regionCode"] = regionCode
    
    if(language != ''):
        request_parameters["relevanceLanguage"] = language
    
    # Remove None values from request_parameters
    request_parameters = {k: v for k, v in request_parameters.items() if v is not None}

    request = youtube.search().list(**request_parameters)
    response = request.execute()
    result_list = []

    for item in response['items']:
        if Type == 'channel':
            result_list.append(item['id']['channelId'])
        elif Type == 'video':
            result_list.append(item['id']['videoId'])
        elif Type == 'playlist':
            playlist_id = item['id'].get('playlistId')
            if playlist_id is not None:
                result_list.append(playlist_id)
            else:
                print(f"Warning: 'playlistId' not found for item {item['snippet']}")

    return result_list

import re
import re

def extractIdFromUrl(url):

    # Regular expressions for different YouTube URL formats
    channel_id_pattern = r'youtube\.com/channel/([a-zA-Z0-9_-]+)'
    channel_name_pattern = r'youtube\.com/(c/|user/|@)([a-zA-Z0-9_-]+)'
    playlist_pattern = r'youtube\.com/playlist\?list=([a-zA-Z0-9_-]+)'
    video_pattern = r'youtube\.com/watch\?v=([a-zA-Z0-9_-]+)'
    short_video_pattern = r'youtu\.be/([a-zA-Z0-9_-]+)'
    youtube_shorts_pattern = r'youtube\.com/shorts/([a-zA-Z0-9_-]+)'

    channel_id_match = re.search(channel_id_pattern, url)
    channel_name_match = re.search(channel_name_pattern, url)
    playlist_match = re.search(playlist_pattern, url)
    video_match = re.search(video_pattern, url)
    short_video_match = re.search(short_video_pattern, url)
    youtube_shorts_match = re.search(youtube_shorts_pattern, url)

    if channel_id_match:
        return channel_id_match.group(1)
    elif channel_name_match:
        return get_channel_id_from_custom_url(url)
    elif playlist_match:
        return playlist_match.group(1)
    elif video_match:
        return video_match.group(1)
    elif short_video_match:
        return short_video_match.group(1)
    elif youtube_shorts_match:
        return youtube_shorts_match.group(1)
    else:
        return None

    
import requests
from bs4 import BeautifulSoup

def get_channel_id_from_custom_url(custom_url):
    try:
        response = requests.get(custom_url)
        if response.status_code != 200:
            print("Failed to retrieve the web page")
            return None

        soup = BeautifulSoup(response.content, 'html.parser')
        meta_tag = soup.find("meta", property="og:url")

        if not meta_tag or not meta_tag.get("content"):
            print("Meta tag with channel URL not found")
            return None

        channel_url = meta_tag.get("content")
        channel_id = channel_url.split('/')[-1]  # Extract the ID from the URL
        return channel_id
    except Exception as e:
        print(f'An error occurred: {e}')
        return None



def get_videos_info(entity_id, youtube, entity_type='channel'):
    next_page_token = None
    unique_tags = set()
    category_count = Counter()
    total_likes = 0  # Initialize total likes
    total_views = 0  # Initialize total views
    total_comments = 0
    name = None  # Initialize name variable
    video_data_list = []  # List to store video data
    top_video_info = None  # Initialize top video info
    durations = []
    # Determine the type of request based on entity type
    if entity_type == 'channel':
        playlist_id = youtube.channels().list(id=entity_id, part='contentDetails').execute()['items'][0]['contentDetails']['relatedPlaylists']['uploads']
    elif entity_type == 'playlist':
        playlist_id = entity_id
    else:
        raise ValueError("Invalid entity type. Must be 'channel' or 'playlist'.")

    # Fetch the videos
    while True:
        res = youtube.playlistItems().list(
            playlistId=playlist_id,
            part='contentDetails,snippet',
            maxResults=50,
            pageToken=next_page_token
        ).execute()

        video_ids = [item['contentDetails']['videoId'] for item in res['items']]
        video_data_list.extend(res['items'])
        next_page_token = res.get('nextPageToken')
        if next_page_token is None:
            break

    top_views = 0  # Variable to store the highest number of views

    for video_id in video_ids:
        video_response = youtube.videos().list(part='statistics,contentDetails,snippet', id=video_id).execute()
        if 'items' in video_response and video_response['items']:
            video_details = video_response['items'][0]
            unique_tags.update(video_details['snippet'].get('tags', []))
            category_count[video_details['snippet']['categoryId']] += 1
            likes = int(video_details['statistics'].get('likeCount', 0))
            views = int(video_details['statistics'].get('viewCount', 0))
            comments = int(video_details['statistics'].get('commentCount', 0))
            total_likes += likes
            total_views += views
            duration = parse_duration_to_minutes(video_details['contentDetails'].get('duration'))
            durations.append(duration)
            # Check if this video has the most views
            if views > top_views:
                top_views = views
                top_video_info = {
                    'videoId': video_details['id'],
                    'title': video_details['snippet']['title'],
                    'description': video_details['snippet']['description'],
                    'thumbnail': video_details['snippet']['thumbnails'].get('high', {}).get('url'),
                    'viewsCount': views,
                    'likesCount': likes,
                    'duration': duration
                }

            # Set name if not already set
            if name is None:
                name = video_details['snippet']['channelTitle'] if entity_type == 'channel' else video_details['snippet']['title']




    # Convert category IDs to names
    category_names = {}
    category_ids = list(category_count.keys())
    if category_ids:
        category_response = youtube.videoCategories().list(part='snippet', id=','.join(category_ids)).execute()
        for item in category_response['items']:
            category_names[item['id']] = item['snippet']['title']

    most_used_categories = [(category_names.get(cid, 'Unknown'), count) for cid, count in category_count.most_common()]

    if top_video_info:
        try:
            # Fetch the top 3 comments for the most viewed video
            comments_response = youtube.commentThreads().list(
                part='snippet',
                videoId=top_video_info['videoId'],
                maxResults=5  # Adjust the maxResults if needed
            ).execute()

            comments = [clean_text(comment['snippet']['topLevelComment']['snippet']['textDisplay']) 
                        for comment in comments_response['items']]
            top_comments = comments[:5]

            # Perform sentiment analysis on these comments
            comment_sentiments = analyze_sentiment(comments)

            # Calculate sentiment percentages
            sentiment_counts = Counter(comment_sentiments)
            total_comments = len(comment_sentiments)
            sentiment_percentages = {sentiment: count / total_comments 
                                    for sentiment, count in sentiment_counts.items()}

            top_video_info['topComments'] = top_comments
            top_video_info['commentSentiments'] = sentiment_percentages
        except:
            print(f"Comments are disabled for video ID {top_video_info['videoId']}. Error: {e}")
            # Optional: set comments-related info to None or an empty structure
            top_video_info['topComments'] = []
            top_video_info['commentSentiments'] = {}


    return {
        'name': name,
        'topVideo': top_video_info,
        'durations': durations,
        'uniqueTags': list(unique_tags),
        'mostUsedCategories': most_used_categories,
        'totalLikes': total_likes,
        'totalViews': total_views,
        # 'videoData': video_data_list  # Optionally include detailed video data
    }



def parse_duration_to_minutes(duration):
    pattern = r'P(?:(\d+)D)?T(?:(\d+)H)?(?:(\d+)M)?(?:(\d+)S)?'
    match = re.match(pattern, duration)
    if match:
        days, hours, minutes, seconds = map(lambda x: int(x) if x else 0, match.groups())
        total_minutes = days * 24 * 60 + hours * 60 + minutes + seconds / 60
        return int(total_minutes)
    else:
        return 0

def analyze_sentiment(comments, max_length=512):
    # Load pre-trained model and tokenizer from Hugging Face
    sentiment_pipeline = pipeline("sentiment-analysis")

    # Truncate comments to the maximum length allowed by the model
    truncated_comments = [comment[:max_length] for comment in comments]

    # Analyze sentiment of each comment
    results = sentiment_pipeline(truncated_comments)

    # Extract sentiment labels from results
    sentiments = [result['label'] for result in results]
    return sentiments

def analyze_youtube_entity(entity_id, youtube, entity_type='channel'):
    data_list = []

    # Fetch entity information using the unified function
    videos_info = get_videos_info(entity_id, youtube, entity_type)
    categ = videos_info['mostUsedCategories']
    topTags = videos_info['uniqueTags'][:30]
    name = videos_info['name']
    topVideo = videos_info['topVideo']
    durations = videos_info['durations']
    mostUsedCategories = videos_info['mostUsedCategories']

    if entity_type == 'channel':
        # Fetch additional channel details
        request_entity = youtube.channels().list(part="snippet,statistics", id=entity_id)
        response_entity = request_entity.execute()

        if not response_entity['items']:
            return pd.DataFrame()  # Return an empty DataFrame if the channel is empty

        entity_stat = response_entity['items'][0]['statistics']
        entity_snippet = response_entity['items'][0]['snippet']
        channel_icon_url = entity_snippet['thumbnails'].get('high', {}).get('url')
        channel_url = f"https://www.youtube.com/channel/{entity_id}"


        # Calculate average views from the original viewCount
        total_views = int(entity_stat['viewCount'])
        subscriber_count = entity_stat.get('subscriberCount', 0)

        request_playlists = youtube.playlists().list(part="snippet", channelId=entity_id)
        response_playlists = request_playlists.execute()
        playlist_count = response_playlists.get("pageInfo", {}).get("totalResults", 0)

        data_list.append([
            channel_url,
            name,
            total_views,
            videos_info['totalViews'],
            categ,
            entity_stat['videoCount'],
            subscriber_count,
            topTags,
            mostUsedCategories,
            playlist_count
        ])
        columns = ['Channel URL', 'Name', 'TotalViews', 'TotalLikes', 'Categories', 'Video count', 'Subscriber count', 'Top tags', 'mostUsedCategories', 'Playlist count']
    else:  # For playlist
        # Use averageViews for playlists
        total_views = videos_info['totalViews']
        request_entity = youtube.playlists().list(part="snippet", id=entity_id)
        response_entity = request_entity.execute()

        # Fetch channel subscriber count
        channel_id = response_entity['items'][0]['snippet']['channelId']
        channel_response = youtube.channels().list(part="statistics", id=channel_id).execute()
        channel_stat = channel_response['items'][0]['statistics']
        subscriber_count = channel_stat.get('subscriberCount', 0)

        
        if not response_entity['items']:
            return pd.DataFrame()  # Return an empty DataFrame if the playlist is empty

        playlist_snippet = response_entity['items'][0]['snippet']
        channel_icon_url = playlist_snippet['thumbnails']['high']['url']
        channel_url = f"https://www.youtube.com/playlist?list={entity_id}"

        data_list.append([
            channel_url,
            name,
            total_views,
            videos_info['totalLikes'],
            categ,
            len(videos_info['durations']),
            subscriber_count,
            topTags,
            mostUsedCategories
        ])
        columns = ['Playlist URL', 'Name', 'TotalViews', 'TotalLikes', 'Categories', 'Video count', 'Subscriber count', 'Top tags', 'mostUsedCategories']

    df = pd.DataFrame(data_list, columns=columns)


    return df, topVideo, channel_icon_url, durations

# # Example usage
# entity_id = id['id']  # Can be either a channel ID or a playlist ID
# entity_type = id['type']  # Change to "playlist" to analyze a playlist
# df, topVideo, channel_icon_url, durations = analyze_youtube_entity(entity_id, youtube, entity_type)


def analyse_channels(ids_list, entity_type, youtube): 
    channel_df_list = []  # List to store dataframes for each channel
    top_videos_list = []  # List to store top video data
    channel_icons_list = []  # List to store channel icons
    durations_list = []
    
    for id in ids_list:
        try:
            df, topVideo, channel_icon_url, durations = analyze_youtube_entity(id, youtube, entity_type)
            channel_df_list.append(df)
            top_videos_list.append(topVideo)
            channel_icons_list.append(channel_icon_url)
            durations_list.append(durations)
            
        except Exception as e:
            print(f"Error analyzing channel {id}: {str(e)}")

    # Concatenate the list of DataFrames vertically
    if channel_df_list:
        merged_channel_df = pd.concat(channel_df_list, ignore_index=True)
    else:
        print("No channel DataFrames to concatenate.")
        merged_channel_df = pd.DataFrame()

    # Convert top videos list to DataFrame
    if top_videos_list:
        top_videos_df = pd.DataFrame(top_videos_list)
    else:
        print("No top videos data to concatenate.")
        top_videos_df = pd.DataFrame()

    return merged_channel_df, top_videos_df, channel_icons_list, durations_list


#video analysis utils
def video_analysis(youtube, video_id):
    next_page_token = None
    unique_tags = set()
    category_count = Counter()
    name = None  # Initialize name variable
    
    video_response = youtube.videos().list(part='statistics,contentDetails,snippet', id=video_id).execute()
    if 'items' in video_response and video_response['items']:
        video_details = video_response['items'][0]
        unique_tags.update(video_details['snippet'].get('tags', []))
        category_count[video_details['snippet']['categoryId']] += 1
        likes = int(video_details['statistics'].get('likeCount', 0))
        views = int(video_details['statistics'].get('viewCount', 0))
        duration = parse_duration_to_minutes(video_details['contentDetails'].get('duration'))
        comments = int(video_details['statistics'].get('commentCount', 0))
        video_info = {
                'videoId': video_details['id'],
                'title': video_details['snippet']['title'],
                'description': video_details['snippet']['description'],
                'thumbnail': video_details['snippet']['thumbnails'].get('high', {}).get('url'),
                'viewsCount': views,
                'likesCount': likes,
                'commentCount': comments,
                'duration': duration,
                'unique_tags' :unique_tags,
                
            }

    # Convert category IDs to names
    category_names = {}
    category_ids = list(category_count.keys())
    if category_ids:
        category_response = youtube.videoCategories().list(part='snippet', id=','.join(category_ids)).execute()
        for item in category_response['items']:
            category_names[item['id']] = item['snippet']['title']

    most_used_categories = [(category_names.get(cid, 'Unknown'), count) for cid, count in category_count.most_common()]
    video_info['most_used_categories'] = most_used_categories
    
    if video_info:
        try:
            # Fetch the top 3 comments for the most viewed video
            comments_response = youtube.commentThreads().list(part='snippet', videoId=video_info['videoId']).execute()
            comments = [clean_text(comment['snippet']['topLevelComment']['snippet']['textDisplay']) for comment in comments_response['items']]
        
            video_info['comments'] = comments
        except: 
            video_info['comments'] = []

    return video_info

import pandas as pd



from transformers import pipeline, BertTokenizer

# Load the tokenizer
tokenizer = BertTokenizer.from_pretrained('bhadresh-savani/bert-base-go-emotion')

# Load the emotion analysis pipeline
emotion_pipeline = pipeline('text-classification', model='bhadresh-savani/bert-base-go-emotion')

def analyze_comments_emotions(comments_df):
    
    tokenizer = BertTokenizer.from_pretrained('bhadresh-savani/bert-base-go-emotion')
    model = AutoModelForSequenceClassification.from_pretrained('bhadresh-savani/bert-base-go-emotion')

    # Function to safely analyze emotion (with truncation)
    def get_emotion(texts):
      # Ensure the input is a list of strings
      if isinstance(texts, str):
          texts = [texts]
      elif not isinstance(texts, list) or not all(isinstance(t, str) for t in texts):
          raise ValueError("Input must be a string or a list of strings.")

      # Tokenize and truncate texts
      inputs = tokenizer(texts, padding=True, truncation=True, max_length=512, return_tensors="pt")
      
      # Predict emotions
      with torch.no_grad():
          outputs = model(**inputs)
      logits = outputs.logits
      predictions = torch.argmax(logits, dim=-1)
      
      # Convert predictions to labels
      labels = [model.config.id2label[prediction.item()] for prediction in predictions][0]
      return labels
    
    # Add a new column to the DataFrame for the predicted emotions
    comments_df['emotion'] = comments_df['text'].apply(get_emotion)

    # Calculate the frequency of each emotion
    emotion_counts = comments_df['emotion'].value_counts(normalize=True)
    
    # Get the top 5 emotions
    top_emotions = emotion_counts.nlargest(5).index.tolist()

    # Create a dictionary to store the top comments for each emotion
    top_comments_by_emotion = {emotion: [] for emotion in top_emotions}

    # Get the top 5 comments for each of the top emotions
    for emotion in top_emotions:
        top_comments = comments_df[comments_df['emotion'] == emotion].nlargest(5, 'likeCount')
        top_comments_by_emotion[emotion] = top_comments.to_dict(orient='records')

    return emotion_counts, top_comments_by_emotion

def video_analysis(youtube, video_id):
    unique_tags = set()
    category_count = Counter()

    

    video_response = youtube.videos().list(
        part='statistics,contentDetails,snippet',
        id=video_id
    ).execute()

    if 'items' in video_response and video_response['items']:
        video_details = video_response['items'][0]
        statistics = video_details['statistics']

        unique_tags.update(video_details['snippet'].get('tags', []))
        category_count[video_details['snippet']['categoryId']] += 1
        
        # Construct a dictionary with all the video info
        video_info = {
            'videoId': video_details['id'],
            'title': video_details['snippet']['title'],
            'description': video_details['snippet']['description'],
            'thumbnail': video_details['snippet']['thumbnails'].get('high', {}).get('url'),
            'viewsCount': int(statistics.get('viewCount', 0)),
            'likesCount': int(statistics.get('likeCount', 0)),
            'commentCount': int(statistics.get('commentCount', 0)),
            'duration': parse_duration_to_minutes(video_details['contentDetails'].get('duration')),
            'unique_tags': list(unique_tags),

        }
        
    # Convert category IDs to names
    category_names = {}
    category_ids = list(category_count.keys())
    if category_ids:
        category_response = youtube.videoCategories().list(part='snippet', id=','.join(category_ids)).execute()
        for item in category_response['items']:
            category_names[item['id']] = item['snippet']['title']

    most_used_categories = [(category_names.get(cid, 'Unknown'), count) for cid, count in category_count.most_common()]
    video_info['most_used_categories'] = most_used_categories
    

    # Create a DataFrame from video info
    video_info_df = pd.DataFrame([video_info])

    comments_data = []

    # Fetch all comments
    next_page_token = None
    while True:
        comments_response = youtube.commentThreads().list(
            part='snippet',
            videoId=video_id,
            pageToken=next_page_token,
            maxResults=100  # Adjust the maxResults if needed
        ).execute()

        
        # Extract comment details and add to the list
        for item in comments_response['items']:
            comment = item['snippet']['topLevelComment']['snippet']
            comment_text = comment['textDisplay']
            clean_comment = clean_text(comment_text)
            comments_data.append({
                'commentId': item['snippet']['topLevelComment']['id'],
                'author': comment['authorDisplayName'],
                'text': clean_comment,
                'likeCount': int(comment.get('likeCount', 0)),
                'replyCount': int(item['snippet']['totalReplyCount']),
                'timestamp': comment['publishedAt']
            })

        # Check if there is a next page
        next_page_token = comments_response.get('nextPageToken')
        if not next_page_token:
            break
    if(comments_data != []):
        # Create a DataFrame from comments data
        comments_df = pd.DataFrame(comments_data)
        print(comments_df.columns)
        emotion_counts, top_comments_by_emotion = analyze_comments_emotions(comments_df)
    else:
        emotion_counts = pd.DataFrame()
        top_comments_by_emotion = {}
        comments_df = pd.DataFrame()
    return video_info_df, comments_df, emotion_counts, top_comments_by_emotion

import googleapiclient.discovery
import pandas as pd
import pandas as pd
from datetime import timedelta
import re
from collections import Counter
from transformers import pipeline


# #Maha=   AIzaSyB5Mi7IXiOBEq5f7nk_kIiq-bVZ6m25rwE
# #Qusasa=   AIzaSyBTkp8Z7xgdHMF8y7BBlWqUabqbERDhyFM
# dev="AIzaSyB5Mi7IXiOBEq5f7nk_kIiq-bVZ6m25rwE"
# api_service_name = "youtube"
# api_version = "v3"
# DEVELOPER_KEY = dev

# youtube = googleapiclient.discovery.build(
#     api_service_name, api_version, developerKey='AIzaSyB5Mi7IXiOBEq5f7nk_kIiq-bVZ6m25rwE')

def parse_duration_to_minutes(duration):
    pattern = r'P(?:(\d+)D)?T(?:(\d+)H)?(?:(\d+)M)?(?:(\d+)S)?'
    match = re.match(pattern, duration)
    if match:
        days, hours, minutes, seconds = map(lambda x: int(x) if x else 0, match.groups())
        total_minutes = days * 24 * 60 + hours * 60 + minutes + seconds / 60
        return int(total_minutes)
    else:
        return 0

def analyse_video(youtube, video_id):
    unique_tags = set()
    category_count = Counter()

    

    video_response = youtube.videos().list(
        part='statistics,contentDetails,snippet',
        id=video_id
    ).execute()

    video_details = video_response['items'][0]
    statistics = video_details['statistics']

    unique_tags.update(video_details['snippet'].get('tags', []))
    category_count[video_details['snippet']['categoryId']] += 1
    
    # Construct a dictionary with all the video info
    video_info = {
        'videoId': video_details['id'],
        'title': video_details['snippet']['title'],
        'description': video_details['snippet']['description'],
        'publishedAt': video_details['snippet']['publishedAt'],
        'thumbnail': video_details['snippet']['thumbnails'].get('high', {}).get('url'),
        'viewsCount': int(statistics.get('viewCount', 0)),
        'likesCount': int(statistics.get('likeCount', 0)),
        'commentCount': int(statistics.get('commentCount', 0)),
        'duration': parse_duration_to_minutes(video_details['contentDetails'].get('duration')),
        'unique_tags': list(unique_tags),
        'categoryId': video_details['snippet']['categoryId']
    }
    
    # Convert category IDs to names
    category_names = {}
    category_ids = list(category_count.keys())
    if category_ids:
        category_response = youtube.videoCategories().list(part='snippet', id=','.join(category_ids)).execute()
        for item in category_response['items']:
            category_names[item['id']] = item['snippet']['title']

    most_used_categories = [(category_names.get(cid, 'Unknown'), count) for cid, count in category_count.most_common()]
    video_info['most_used_categories'] = most_used_categories
    

    # Create a DataFrame from video info
    video_info_df = pd.DataFrame([video_info])
    return video_info_df

def analyse_comments_data(youtube, video_id):
    comments_data = []

    try:
        # Fetch all comments
        comments_response = youtube.commentThreads().list(
            part='snippet',
            videoId=video_id,
            maxResults=30  # Adjust the maxResults if needed
        ).execute()

        # Extract comment details and add to the list
        for item in comments_response['items']:
            comment = item['snippet']['topLevelComment']['snippet']
            comment_text = comment['textDisplay']
            clean_comment = clean_text(comment_text)
            comments_data.append({
                'commentId': item['snippet']['topLevelComment']['id'],
                'author': comment['authorDisplayName'],
                'text': clean_comment,
                'likeCount': int(comment.get('likeCount', 0)),
                'replyCount': int(item['snippet']['totalReplyCount']),
                'timestamp': comment['publishedAt']
            })

    except HttpError as e:
        # Handle the case where comments are disabled or other HTTP errors
        print(f"Error fetching comments for video ID {video_id}: {e}")
        # You can return an empty list or a DataFrame, or handle this in another way
        return [{
                'commentId': '',
                'author': '',
                'text': 'No comments',
                'likeCount': 0,
                'replyCount': 0,
                'timestamp': ''
            }]

    # Create a DataFrame from comments data if needed
    # comments_df = pd.DataFrame(comments_data)
    # return comments_df

    return comments_data


def analyze_comments_emotions_for_playlist(comments_df):
    
    tokenizer = BertTokenizer.from_pretrained('bhadresh-savani/bert-base-go-emotion')
    model = AutoModelForSequenceClassification.from_pretrained('bhadresh-savani/bert-base-go-emotion')

    # Function to safely analyze emotion (with truncation)
    def get_emotion(texts):
      
      # Ensure the input is a list of strings
      if isinstance(texts, str):
          texts = [texts]
      elif not isinstance(texts, list) or not all(isinstance(t, str) for t in texts):
          raise ValueError("Input must be a string or a list of strings.")

      # Tokenize and truncate texts
      inputs = tokenizer(texts, padding=True, truncation=True, max_length=512, return_tensors="pt")
      
      # Predict emotions
      with torch.no_grad():
          outputs = model(**inputs)
      logits = outputs.logits
      predictions = torch.argmax(logits, dim=-1)
      
      # Convert predictions to labels
      label = [model.config.id2label[prediction.item()] for prediction in predictions][0]
      return label
    
    # Add a new column to the DataFrame for the predicted emotions
    comments_df['emotion'] = comments_df['text'].apply(get_emotion)

    # Calculate the frequency of each emotion
    emotion_counts = comments_df['emotion'].value_counts(normalize=True)
    # Get the top 5 emotions
    top_emotions = emotion_counts.nlargest(5).index.tolist()

    # Create a dictionary to store the top comments for each emotion
    top_comments_by_emotion = {emotion: [] for emotion in top_emotions}

    # Get the top 5 comments for each of the top emotions
    for emotion in top_emotions:
        top_comments = comments_df[comments_df['emotion'] == emotion].nlargest(1, 'likeCount')
        top_comments_by_emotion[emotion] = top_comments['text'].iloc[0]
        

    return emotion_counts, top_comments_by_emotion


from datetime import datetime

def calculate_engagement_score(video_info, current_time):
    days_since_published = (current_time - datetime.fromisoformat(video_info['publishedAt'][:-1])).days
    if days_since_published == 0:
        days_since_published = 1  # To avoid division by zero
    average_counts = (video_info['viewsCount'] + video_info['likesCount'] + video_info['commentCount']) / 3
    engagement_score = average_counts / days_since_published
    return engagement_score


def analyze_playlist(youtube, playlist_id):
    # Fetch playlist details
    playlist_response = youtube.playlists().list(
        part="snippet,contentDetails",
        id=playlist_id
    ).execute()

    playlist_details = playlist_response['items'][0]
    
    # Fetch videos in the playlist
    next_page_token = None
    videos_info_list = []
    while True:
        playlist_videos_response = youtube.playlistItems().list(
            part="snippet,contentDetails",
            playlistId=playlist_id,
            maxResults=50,
            pageToken=next_page_token
        ).execute()

        for item in playlist_videos_response['items']:
            try:
                video_id = item['contentDetails']['videoId']
                video_info_df = analyse_video(youtube, video_id)
                video_info = video_info_df.iloc[0]
                video_info['engagementScore'] = calculate_engagement_score(video_info, datetime.now())
                videos_info_list.append(video_info)
            except:
                continue
        next_page_token = playlist_videos_response.get('nextPageToken')
        if not next_page_token:
            break

    # Create a DataFrame from all videos info
    all_videos_info_df = pd.DataFrame(videos_info_list)
    
    

    # Sort by engagement score instead of views
    top_3_videos = all_videos_info_df.nlargest(3, 'engagementScore')
    worst_3_videos = all_videos_info_df.nsmallest(3, 'engagementScore')


    top_3_comments = []
    for _, row in top_3_videos.iterrows():
        video_comments = analyse_comments_data(youtube, row['videoId'])
        if video_comments != []:
            top_3_comments.extend(video_comments)

    # Collect comments for worst 3 videos
    worst_3_comments = []
    for _, row in worst_3_videos.iterrows():
        video_comments = analyse_comments_data(youtube, row['videoId'])
        if video_comments != []:
            worst_3_comments.extend(video_comments)

    if top_3_comments != []:
    # Convert comments lists to DataFrames
        top_3_comments_df = pd.DataFrame(top_3_comments)
        top_3_comments_analysis = analyze_comments_emotions_for_playlist(top_3_comments_df)
    else:
        top_3_comments_df = []
        top_3_comments_analysis = []
        
    if worst_3_comments != []:
        worst_3_comments_df = pd.DataFrame(worst_3_comments)
        worst_3_comments_analysis = analyze_comments_emotions_for_playlist(worst_3_comments_df)
    else:
        worst_3_comments_df = []
        worst_3_comments_analysis = []
    # Calculate aggregated data
    total_views = all_videos_info_df['viewsCount'].sum()
    total_likes = all_videos_info_df['likesCount'].sum()
    total_comments = all_videos_info_df['commentCount'].sum()
    average_duration = all_videos_info_df['duration'].mean()

    # Aggregate unique tags
    unique_tags = set()
    for tags in all_videos_info_df['unique_tags']:
        unique_tags.update(tags)

    # Update playlist_info with aggregated data
    playlist_info = {
        'playlistId': playlist_id,
        'title': playlist_details['snippet']['title'],
        'description': playlist_details['snippet']['description'],
        'thumbnail': playlist_details['snippet']['thumbnails'].get('high', {}).get('url'),
        'channelName': playlist_details['snippet']['channelTitle'],
        'publishedAt': playlist_details['snippet']['publishedAt'],
        'videoCount': playlist_details['contentDetails']['itemCount'],
        'totalViews': total_views,
        'totalLikes': total_likes,
        'totalComments': total_comments,
        'average_duration': average_duration,
        'uniqueTags': list(unique_tags)
    }

    # Create a DataFrame from playlist info
    playlist_info_df = pd.DataFrame([playlist_info])

    return playlist_info_df, all_videos_info_df, top_3_videos, worst_3_videos, top_3_comments_analysis, worst_3_comments_analysis




def analyze_channel(youtube, channel_id):
    # Fetch channel details
    channel_response = youtube.channels().list(
        part="snippet,contentDetails,statistics",
        id=channel_id
    ).execute()

    channel_details = channel_response['items'][0]

    # Extract channel info
    channel_info = {
        'Channel Name': channel_details['snippet']['title'],
        'description': channel_details['snippet']['description'],
        'thumbnail': channel_details['snippet']['thumbnails'].get('high', {}).get('url'),
        'viewCount': channel_details['statistics']['viewCount'],
        'videoCount': channel_details['statistics']['videoCount'],
        'subscriberCount': channel_details['statistics']['subscriberCount'],
        'publishedAt': channel_details['snippet']['publishedAt'],

    }

    # Fetch and analyze videos in the channel
    all_videos_info = []
    next_page_token = None
    while True:
        video_response = youtube.search().list(
            part="snippet",
            channelId=channel_id,
            maxResults=50,
            pageToken=next_page_token,
            type="video"
        ).execute()
        
        category_count = Counter()
        for item in video_response['items']:
            video_id = item['id']['videoId']
            video_info = analyse_video(youtube, video_id)
            video_info = video_info.iloc[0]
            video_info['engagementScore'] = calculate_engagement_score(video_info, datetime.now())
            category_id = video_info.get('categoryId')
            if category_id:
                category_count[category_id] += 1
                
            all_videos_info.append(video_info)

        next_page_token = video_response.get('nextPageToken')
        if not next_page_token:
            break

    all_videos_df = pd.DataFrame(all_videos_info)
    # Convert category IDs to names and find the most used categories
    category_names = {}
    category_ids = list(category_count.keys())
    if category_ids:
        category_response = youtube.videoCategories().list(
            part='snippet',
            id=','.join(category_ids)
        ).execute()

        for item in category_response['items']:
            category_names[item['id']] = item['snippet']['title']

    most_used_categories = [(category_names.get(cid, 'Unknown'), count) 
                            for cid, count in category_count.most_common()]

    # Add most used categories to channel info
    channel_info['mostUsedCategories'] = most_used_categories
    
    total_likes = all_videos_df['likesCount'].sum()
    total_comments = all_videos_df['commentCount'].sum()
    average_duration = all_videos_df['duration'].mean()

    channel_info['likesCount'] = total_likes
    channel_info['commentCount'] = total_comments
    channel_info['average_duration'] = average_duration
    
    # Sort by engagement score instead of views
    top_3_videos = all_videos_df.nlargest(3, 'engagementScore')
    worst_3_videos = all_videos_df.nsmallest(3, 'engagementScore')


    top_3_comments = []
    for _, row in top_3_videos.iterrows():
        video_comments = analyse_comments_data(youtube, row['videoId'])
        if video_comments != []:
            top_3_comments.extend(video_comments)

    # Collect comments for worst 3 videos
    worst_3_comments = []
    for _, row in worst_3_videos.iterrows():
        video_comments = analyse_comments_data(youtube, row['videoId'])
        if video_comments != []:
            worst_3_comments.extend(video_comments)

    if top_3_comments != []:
    # Convert comments lists to DataFrames
        top_3_comments_df = pd.DataFrame(top_3_comments)
        top_3_comments_analysis = analyze_comments_emotions_for_playlist(top_3_comments_df)
    else:
        top_3_comments_df = []
        top_3_comments_analysis = []
        
    if worst_3_comments != []:
        worst_3_comments_df = pd.DataFrame(worst_3_comments)
        worst_3_comments_analysis = analyze_comments_emotions_for_playlist(worst_3_comments_df)
    else:
        worst_3_comments_df = []
        worst_3_comments_analysis = []
        
    # Aggregate unique tags
    unique_tags = set()
    for tags in all_videos_df['unique_tags']:
        unique_tags.update(tags)
    
    channel_info['uniqueTags'] = unique_tags
    
    all_playlists_info = []
    next_page_token = None
    while True:
        playlist_response = youtube.playlists().list(
            part="snippet,contentDetails",
            channelId=channel_id,
            maxResults=50,
            pageToken=next_page_token
        ).execute()

        for item in playlist_response['items']:
            playlist_id = item['id']
            playlist_details = item['snippet']
            content_details = item['contentDetails']

            playlist_info = {
                'playlistId': playlist_id,
                'title': playlist_details['title'],
                'description': playlist_details['description'],
                'thumbnail': playlist_details['thumbnails'].get('high', {}).get('url'),
                'publishedAt': playlist_details['publishedAt'],
                'videoCount': content_details['itemCount'],
            }
            all_playlists_info.append(playlist_info)

        next_page_token = playlist_response.get('nextPageToken')
        if not next_page_token:
            break
        
    channel_info['PlaylistCount'] = len(all_playlists_info)
    
    all_playlists_df = pd.DataFrame(all_playlists_info)
    channel_df = pd.DataFrame([channel_info])

    return channel_df, all_videos_df, all_playlists_df, top_3_videos, worst_3_videos, top_3_comments_analysis, worst_3_comments_analysis



from pytube import YouTube
import subprocess
import os

def download_audio_from_youtube(url, output_dir):
    yt = YouTube(url)
    audio_stream = yt.streams.filter(only_audio=True).first()

    # Download and save the audio stream
    download_path = audio_stream.download(output_path=output_dir)

    # Construct the mp3 filename
    mp3_filename = os.path.join(output_dir, os.path.splitext(os.path.basename(download_path))[0] + ".mp3")
    
    # Use FFmpeg to convert the downloaded file to mp3
    subprocess.run(['ffmpeg', '-i', download_path, mp3_filename])

    # Optionally, delete the original download
    os.remove(download_path)

    return mp3_filename

import whisper

def transcribe_youtube_video(audio_file):
    model = whisper.load_model("tiny")
    result = model.transcribe(audio_file)
    return result["text"]

# youtube_downloader.py

import openai

def summarize_youtube_video(transcript, openai_api_key):
    openai.api_key = openai_api_key

    response = openai.Completion.create(
      engine="text-davinci-003",
      prompt=f"Provide a summary for the following text:\n\n{transcript}",
      max_tokens=150
    )

    return response.choices[0].text.strip()

# youtube_downloader.py

from docx import Document
import os

def create_word_document(transcript, summary, filename, output_dir):
    doc = Document()
    doc.add_heading('YouTube Video Analysis', 0)

    doc.add_heading('Transcript:', level=1)
    doc.add_paragraph(transcript)

    doc.add_heading('Summary:', level=1)
    doc.add_paragraph(summary)

    # Save the document
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    file_path = os.path.join(output_dir, f"{filename}.docx")
    doc.save(file_path)

    return file_path

def topic_analysis(youtube, query, orderBy='relevance', regionCode='', language='', max_results=100):

    one_year_ago = datetime.now() - timedelta(days=365)
    one_year_ago_str = one_year_ago.strftime('%Y-%m-%dT%H:%M:%SZ')
    
    kw_model = KeyBERT(model='all-MiniLM-L6-v2')
    descri = ''
    
    
    request_parameters = {
        "part": "snippet",
        "order": orderBy,
        "q": query,
        "type": 'videos',
        "maxResults": max_results,
        "publishedAfter": one_year_ago_str
    }

    if(regionCode != ''):
        request_parameters["regionCode"] = regionCode
    
    if(language != ''):
        request_parameters["relevanceLanguage"] = language

    # Step 1: Search for top videos based on the query
    search_response = youtube.search().list(**request_parameters).execute()

    videos = []
    channels = {}
    all_tags = []
    unique_tags = set()
    
    category_count = Counter()

    for item in search_response.get('items', []):
        video_id = item['id']['videoId']
        channel_id = item['snippet']['channelId']

        # Fetch additional video details
        video_response = youtube.videos().list(
            part='snippet,contentDetails,statistics',
            id=video_id
        ).execute()

        video_details = video_response['items'][0]
        statistics = video_details['statistics']

        unique_tags.update(video_details['snippet'].get('tags', []))
        category_count[video_details['snippet']['categoryId']] += 1

        video_info = {
            'videoId': video_details['id'],
            'title': video_details['snippet']['title'],
            'description': video_details['snippet']['description'],
            'publishedAt': video_details['snippet']['publishedAt'],
            'thumbnail': video_details['snippet']['thumbnails'].get('high', {}).get('url'),
            'viewsCount': int(statistics.get('viewCount', 0)),
            'likesCount': int(statistics.get('likeCount', 0)),
            'commentCount': int(statistics.get('commentCount', 0)),
            'duration': parse_duration_to_minutes(video_details['contentDetails'].get('duration')),
            'unique_tags': list(unique_tags),
            'categoryId': video_details['snippet']['categoryId']

        }
        
        title = video_info['title']
        description = video_info['description']
        descri = descri + str(title) + ', ' + str(description) + ', '
        
        category_names = {}
        category_ids = list(category_count.keys())
        if category_ids:
            category_response = youtube.videoCategories().list(part='snippet', id=','.join(category_ids)).execute()
            for item in category_response['items']:
                category_names[item['id']] = item['snippet']['title']

        most_used_categories = [(category_names.get(cid, 'Unknown'), count) for cid, count in category_count.most_common()]
        video_info['most_used_categories'] = most_used_categories

        video_info['engagementScore'] = calculate_engagement_score(video_info, datetime.now())

        videos.append(video_info)

        
        # Track channel activity
        channels[channel_id] = channels.get(channel_id, 0) + 1

    # Step 2: Fetch the most active channels
    active_channels = sorted(channels.items(), key=lambda x: x[1], reverse=True)[:5]
    channels = []
    for channel_id, _ in active_channels:
        channel_df , topVideo, channel_icon_url, durations = analyze_youtube_entity(channel_id, youtube)
        channel_df['channel_icon_url'] = channel_icon_url
        channels.append(channel_df)

    # Convert results to DataFrames
    videos_df = pd.DataFrame(videos)
    top_5_videos = videos_df.nlargest(5, 'engagementScore')
    
    top_5_comments = []
    for _, row in top_5_videos.iterrows():
        try:
            # Attempt to analyze comments data
            video_comments = analyse_comments_data(youtube, row['videoId'])
            top_5_comments.extend(video_comments)
        except HttpError as e:
            # Handle the case where comments are disabled
            # You can log this error, skip the video, or handle it in another way
            print(f"Comments are disabled for video ID {row['videoId']}. Error: {e}")
            continue  # Skip this video and continue with the next one

    top_5_comments_df = pd.DataFrame(top_5_comments)
    top_5_comments_analysis = analyze_comments_emotions_for_playlist(top_5_comments_df)
        
    channels_df = pd.concat(channels, ignore_index=True)

    

   
    keybert_keywords = kw_model.extract_keywords(descri, keyphrase_ngram_range=(1, 2), top_n=10)

    return videos_df, channels_df, top_5_videos, top_5_comments_df, top_5_comments_analysis, keybert_keywords




def get_realted_videos(youtube, video_id, order='relevance', region_code='', language='', number_of_videos=10):
    # Step 1: Fetch the video's details
    video_details_request = youtube.videos().list(
        part="snippet",
        id=video_id
    )
    video_details_response = video_details_request.execute()

    if not video_details_response['items']:
        return pd.DataFrame()  # Return empty DataFrame if no details are found

    video_details = video_details_response['items'][0]
    tags = video_details['snippet'].get('tags', [])[:5]  # Get top 5 tags
    category_id = video_details['snippet']['categoryId']

    # Calculate the number of videos to fetch per tag

    # Step 2: Use tags to find related videos
    related_videos = []

    print(tags)
    if tags != []:
        videos_per_tag = max(1, number_of_videos // len(tags)) if tags else number_of_videos

        for tag in tags:
            related_videos.extend(
                get_videos(youtube, tag, category_id, order, region_code, language, videos_per_tag)
            )

        # If tags are less than needed, get more videos with the last tag
        if len(tags) < number_of_videos:
            additional_videos_needed = number_of_videos - len(related_videos)
            related_videos.extend(
                get_videos(youtube, tags[-1], category_id, order, region_code, language, additional_videos_needed)
            )
    else:
        keywords = keyword_extraction(video_id)
        videos_per_tag = max(1, number_of_videos // len(keywords)) if keywords else number_of_videos
        for keyword in keywords:
            word = keyword[0]
            related_videos.extend(
                get_videos(youtube, word, category_id, order, region_code, language, videos_per_tag)
            )

        # If tags are less than needed, get more videos with the last tag
        if len(keywords) < number_of_videos:
            additional_videos_needed = number_of_videos - len(related_videos)
            related_videos.extend(
                get_videos(youtube, keywords[-1][0], category_id, order, region_code, language, additional_videos_needed)
            )
    

    # Convert to DataFrame and drop duplicates
    related_videos_df = pd.DataFrame(related_videos).drop_duplicates(subset=['Id'])

    return related_videos_df
    
    
def get_videos(youtube, search_term, video_category_id='none', order='relevance', region_code='', language='', number_of_videos=10):
    # Initialize variables
    related_videos = []
    total_results = 0
    nextPageToken = None
    category_names = {}

    while total_results < number_of_videos:
        # Adjust max_results based on remaining videos needed
        current_max_results = min(100, number_of_videos - total_results)

        search_params = {
            "part": "snippet",
            "q": search_term,
            "type": "video",
            "maxResults": current_max_results,
            "order": order,
            "pageToken": nextPageToken
        }

        # Validate region_code and language
        if region_code and region_code.isalpha() and len(region_code) == 2:
            search_params["regionCode"] = region_code
            
        if language and re.match("^[a-zA-Z-]+$", language):
            search_params["relevanceLanguage"] = language
        
        if video_category_id != 'none':
            search_params["videoCategoryId"] = video_category_id

        search_response = youtube.search().list(**search_params).execute()
        video_ids = [search_result["id"]["videoId"] for search_result in search_response.get("items", [])]

        # Get detailed information for each video
        videos_response = youtube.videos().list(
            part="contentDetails,snippet,statistics",
            id=",".join(video_ids)
        ).execute()

        # Fetch category names if not already fetched
        if not category_names:
            category_response = youtube.videoCategories().list(part="snippet", regionCode=region_code).execute()
            for item in category_response.get("items", []):
                category_names[item["id"]] = item["snippet"]["title"]

        for video in videos_response.get("items", []):
            stats = video["statistics"]
            video_details = {
                "Title": video["snippet"]["title"],
                "Id": video["id"],
                "URL": f"https://www.youtube.com/watch?v={video['id']}",
                "Channel": video["snippet"].get("channelTitle", "Unknown Channel"),
                "Thumbnial": video["snippet"]['thumbnails'].get('high', {}).get('url'),
                "Description": video["snippet"]["description"],
                "Tags": video["snippet"].get("tags", []),
                "Category": category_names.get(video["snippet"]["categoryId"], "Unknown"),
                "Views": stats.get("viewCount", "0"),
                "Likes": stats.get("likeCount", "0"),
                "Dislikes": stats.get("dislikeCount", "0"),
                "Comments": stats.get("commentCount", "0"),
                "Duration": parse_duration_to_minutes(video['contentDetails'].get('duration'))
            }
            related_videos.append(video_details)

        total_results += len(search_response.get("items", []))

        # Check if there are more results and update the nextPageToken
        nextPageToken = search_response.get("nextPageToken")
        if not nextPageToken:
            break
        

    return related_videos


from keybert import KeyBERT

def keyword_extraction(videoid):
    kw_model = KeyBERT(model='all-MiniLM-L6-v2')
    anl = []

    request = youtube.videos().list(
        part="snippet",
        id=videoid,
        regionCode="US"
    )
    response = request.execute()

    for item in response.get('items', []):
        snippet = item.get('snippet', {})
        video_id = item.get('id')
        title = snippet.get('title', '')
        description = snippet.get('description', '')
        anl.append([video_id, title, description])

    descri = str(title) + '\n' + str(description)
    keywords = kw_model.extract_keywords(descri, keyphrase_ngram_range=(1, 1))
    return keywords

import re
import html

def clean_text(input_text):
    # Remove HTML tags and extra whitespaces
    cleaned_text = re.sub(r'<.*?>', '', input_text)
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()

    # Decode HTML entities
    cleaned_text = html.unescape(cleaned_text)

    return cleaned_text